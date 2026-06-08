"""
QPSS 脚本フォーマッタ エンジン
YAMLプリセットに基づいて脚本テキストをWord文書(.docx)に変換する。
"""

import io
import os
import re
from pathlib import Path

import yaml
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

# ページサイズ定義 (幅, 高さ) in mm - portrait基準
PAGE_SIZES = {
    "A4": (210, 297),
    "B5": (182, 257),
    "Letter": (216, 279),
}

PRESETS_DIR = Path(__file__).parent / "presets"


def load_preset(preset_name_or_path):
    """プリセットYAMLを読み込む。名前またはパスを受け付ける。"""
    if os.path.isfile(preset_name_or_path):
        path = Path(preset_name_or_path)
    else:
        path = PRESETS_DIR / f"{preset_name_or_path}.yaml"
        if not path.exists():
            raise FileNotFoundError(f"プリセットが見つかりません: {preset_name_or_path}")
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def list_presets():
    """利用可能なプリセット一覧を返す。"""
    presets = []
    if not PRESETS_DIR.exists():
        return presets
    for p in sorted(PRESETS_DIR.glob("*.yaml")):
        if p.name.startswith("_"):
            continue
        try:
            with open(p, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f)
            presets.append({
                "file": p.name,
                "name": data.get("name", p.stem),
                "description": data.get("description", ""),
            })
        except Exception:
            continue
    return presets


def extract_text(file_path_or_bytes, filename=None):
    """各種ファイル形式からテキストを抽出する。

    Args:
        file_path_or_bytes: ファイルパス(str/Path)またはバイトデータ
        filename: バイトデータの場合のファイル名（拡張子判定用）
    """
    if isinstance(file_path_or_bytes, (str, Path)):
        path = Path(file_path_or_bytes)
        ext = path.suffix.lower()
        with open(path, "rb") as f:
            data = f.read()
    else:
        data = file_path_or_bytes
        ext = Path(filename).suffix.lower() if filename else ""

    if ext == ".txt":
        for encoding in ["utf-8", "shift_jis", "cp932", "euc-jp"]:
            try:
                return data.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                continue
        return data.decode("utf-8", errors="replace")

    elif ext == ".docx":
        try:
            import mammoth
            result = mammoth.extract_raw_text(io.BytesIO(data))
            return result.value
        except ImportError:
            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs)

    elif ext == ".rtf":
        try:
            from striprtf.striprtf import rtf_to_text
            text = data.decode("utf-8", errors="replace")
            return rtf_to_text(text)
        except ImportError:
            raise ImportError("RTFファイルの読み込みには striprtf パッケージが必要です")

    elif ext == ".md":
        for encoding in ["utf-8", "shift_jis", "cp932"]:
            try:
                text = data.decode(encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        else:
            text = data.decode("utf-8", errors="replace")
        # Markdownの装飾記号を除去して素のテキストにする
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)  # 見出し
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # 太字
        text = re.sub(r"\*(.+?)\*", r"\1", text)  # 斜体
        text = re.sub(r"^[-*+]\s+", "", text, flags=re.MULTILINE)  # リスト
        return text

    else:
        # 不明な形式はテキストとして試行
        for encoding in ["utf-8", "shift_jis", "cp932"]:
            try:
                return data.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                continue
        return data.decode("utf-8", errors="replace")


class ScenarioFormatter:
    """YAMLプリセットに基づいて脚本をdocxに変換するフォーマッタ。"""

    def __init__(self, preset):
        """
        Args:
            preset: プリセット辞書（YAMLから読み込んだもの）またはプリセット名/パス
        """
        if isinstance(preset, str):
            self.config = load_preset(preset)
        else:
            self.config = preset

    def is_togaki(self, line):
        """ト書き行かどうかを判定する。除外パターンに該当しない行がト書き。"""
        stripped = line.strip()
        if not stripped:
            return False

        togaki_config = self.config.get("togaki", {})
        if not togaki_config.get("auto_detect", True):
            return False

        patterns = togaki_config.get("exclude_patterns", [])
        for pat in patterns:
            pat_type = pat.get("type", "")
            value = pat.get("value", "")

            if pat_type == "contains" and value in line:
                return False
            elif pat_type == "starts_with" and (
                line.startswith(value) or stripped.startswith(value)
            ):
                # 全角スペース等の空白で始まるパターン（セリフ続き行など）は
                # line.strip() で先頭空白が除去されると一致しなくなるため、
                # 元の行頭(line)でも判定する。
                return False
            elif pat_type == "exact" and stripped == value:
                return False
            elif pat_type == "regex":
                if re.search(value, stripped):
                    return False
            elif pat_type == "char_name":
                max_pos = pat.get("max_paren_pos", 6)
                if "（" in stripped and "）" in stripped and stripped[0] != "（":
                    paren_pos = stripped.index("（")
                    if paren_pos <= max_pos:
                        return False

        return True

    def _add_page_number(self, section, position, font_size):
        """フッター/ヘッダーにページ番号を追加する。"""
        if "top" in position:
            container = section.header
        else:
            container = section.footer

        container.is_linked_to_previous = False
        paragraph = container.paragraphs[0]

        # 揃え設定
        pPr = paragraph._element.get_or_add_pPr()
        jc = OxmlElement("w:jc")
        if "center" in position:
            jc.set(qn("w:val"), "center")
        else:
            jc.set(qn("w:val"), "right")
        pPr.append(jc)

        # PAGE フィールド挿入
        font_name = self.config.get("font", {}).get("name", "MS Gothic")
        run = paragraph.add_run()
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run.element.append(fldChar_begin)

        run2 = paragraph.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run2.element.append(instrText)

        run3 = paragraph.add_run()
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run3.element.append(fldChar_end)

    def create_docx(self, text, writer_name="", output_path=None):
        """プリセット設定に基づいてdocxを生成する。

        Args:
            text: 脚本テキスト
            writer_name: 執筆者名
            output_path: 出力ファイルパス。Noneの場合はBytesIOを返す。

        Returns:
            output_pathが指定された場合はNone、それ以外はBytesIOオブジェクト
        """
        doc = Document()
        section = doc.sections[-1]

        # --- ページ設定 ---
        page_config = self.config.get("page", {})
        size_name = page_config.get("size", "A4")
        w_mm, h_mm = PAGE_SIZES.get(size_name, PAGE_SIZES["A4"])

        orientation = page_config.get("orientation", "landscape")
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Mm(h_mm)
            section.page_height = Mm(w_mm)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Mm(w_mm)
            section.page_height = Mm(h_mm)

        # 縦書き設定
        text_direction = page_config.get("text_direction", "vertical")
        if text_direction == "vertical":
            sectPr = section._sectPr
            textDir = OxmlElement("w:textDirection")
            textDir.set(qn("w:val"), "tbRl")
            sectPr.append(textDir)

        # --- フォント・段落スタイル ---
        font_config = self.config.get("font", {})
        font_name = font_config.get("name", "MS Gothic")
        font_size = font_config.get("size", 12)

        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

        para_config = self.config.get("paragraph", {})
        pf = style.paragraph_format
        indent_before = para_config.get("indent_before_mm", 6)
        indent_after = para_config.get("indent_after_mm", 0)
        # mm → pt 変換 (1mm ≈ 2.83465pt)
        pf.left_indent = Pt(indent_before * 2.83465)
        pf.right_indent = Pt(indent_after * 2.83465)
        pf.space_before = Pt(para_config.get("space_before_pt", 0))
        pf.space_after = Pt(para_config.get("space_after_pt", 0))

        # --- ページ番号 ---
        hf_config = self.config.get("header_footer", {})
        if hf_config.get("page_number", True):
            position = hf_config.get("page_number_position", "right_bottom")
            pn_font_size = hf_config.get("page_number_font_size", 10)
            self._add_page_number(section, position, pn_font_size)

        # --- テキスト流し込み ---
        togaki_tabs = self.config.get("togaki", {}).get("indent_tabs", 2)
        meta_config = self.config.get("meta", {})
        writer_label = meta_config.get("writer_label", "脚本")
        writer_position = meta_config.get("writer_position", "after_title")

        # 行間設定
        line_spacing = self.config.get("line_spacing", {})
        blank_between_dialogue = line_spacing.get("between_dialogue", 0)
        blank_between_togaki_dialogue = line_spacing.get("between_togaki_and_dialogue", 1)

        lines = text.split("\n")
        title_done = False

        # 空行を除去してコンテンツ行のみ抽出し、行間を再構成する
        content_lines = []
        for line in lines:
            # 執筆者名の挿入（タイトル直後）
            if not title_done and line.strip() and writer_name and writer_position == "after_title":
                content_lines.append(("title", line))
                content_lines.append(("meta", f"{writer_label}：{writer_name}"))
                title_done = True
                continue

            if not title_done and line.strip():
                title_done = True

            if not line.strip():
                continue

            is_tog = self.is_togaki(line)
            if is_tog:
                content_lines.append(("togaki", line))
            else:
                content_lines.append(("dialogue", line))

        # コンテンツ行を出力し、行種の変化に応じて空行を挿入
        for idx, (line_type, line) in enumerate(content_lines):
            # 前の行との間に空行を挿入
            if idx > 0:
                prev_type = content_lines[idx - 1][0]
                blanks = 0
                if prev_type in ("title", "meta"):
                    blanks = 0
                elif line_type == "dialogue" and prev_type == "dialogue":
                    blanks = blank_between_dialogue
                else:
                    # ト書き⇔セリフ、ト書き⇔ト書き の切り替わり
                    blanks = blank_between_togaki_dialogue
                for _ in range(blanks):
                    doc.add_paragraph("")

            # ト書き → TABインデント付与
            if line_type == "togaki":
                line = "\t" * togaki_tabs + line

            p = doc.add_paragraph()
            run = p.add_run(line)
            if line_type in ("title", "meta"):
                run.bold = True

        # --- 保存 ---
        if output_path:
            doc.save(output_path)
            return None
        else:
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf


def save_preset(config, filename):
    """プリセット設定をYAMLファイルとして保存する。"""
    path = PRESETS_DIR / filename
    with open(path, "w", encoding="utf-8") as f:
        yaml.dump(config, f, allow_unicode=True, default_flow_style=False, sort_keys=False)
    return path


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="QPSS 脚本フォーマッタ")
    parser.add_argument("--preset", default="default", help="プリセット名またはパス")
    parser.add_argument("--input", "-i", help="入力ファイルパス")
    parser.add_argument("--output", "-o", help="出力ファイルパス")
    parser.add_argument("--writer", "-w", default="", help="執筆者名")
    parser.add_argument("--list-presets", action="store_true", help="プリセット一覧を表示")

    args = parser.parse_args()

    if args.list_presets:
        print("利用可能なプリセット:")
        print("-" * 50)
        for p in list_presets():
            print(f"  {p['file']:<25} {p['name']}")
            if p["description"]:
                print(f"  {'':25} {p['description']}")
            print()
    elif args.input:
        text = extract_text(args.input)
        formatter = ScenarioFormatter(args.preset)
        out = args.output or Path(args.input).stem + "_formatted.docx"
        formatter.create_docx(text, writer_name=args.writer, output_path=out)
        print(f"出力完了: {out}")
    else:
        parser.print_help()
