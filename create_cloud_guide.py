"""Streamlit Cloud版 使い方ガイドのdocxを生成するスクリプト"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


APP_URL = "https://qpss-scenario-formatter-fdrkrqerpcxwyd3wxoslmq.streamlit.app/"


def add_hyperlink(paragraph, text, url):
    """段落にハイパーリンクを追加"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1A56DB")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rPr.append(sz)
    run.append(rPr)
    run.text = text
    hyperlink.append(run)
    paragraph._element.append(hyperlink)
    return paragraph


def add_styled_paragraph(doc, text, bold=False, font_size=None,
                         color=None, alignment=None, space_after=None, space_before=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def create_guide():
    doc = Document()

    # デフォルトスタイル
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Yu Gothic"
        h.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        if level == 1:
            h.font.size = Pt(18)
            h.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        elif level == 2:
            h.font.size = Pt(14)
            h.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        elif level == 3:
            h.font.size = Pt(12)
            h.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # =============================================
    # タイトル
    # =============================================
    add_styled_paragraph(doc, "", space_after=20)

    add_styled_paragraph(
        doc, "QPSS 脚本フォーマッタ",
        bold=True, font_size=26, color=(0x1A, 0x56, 0xDB),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8
    )
    add_styled_paragraph(
        doc, "使い方ガイド（Web版）",
        font_size=16, color=(0x66, 0x66, 0x66),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
    )
    add_styled_paragraph(
        doc, "Version 1.0  |  2026年4月",
        font_size=10, color=(0x99, 0x99, 0x99),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30
    )

    # =============================================
    # アプリURL
    # =============================================
    doc.add_heading("アプリのリンク", level=1)

    p = doc.add_paragraph("以下のURLをクリックするだけで、すぐに使えます。インストール不要です。")
    p.paragraph_format.space_after = Pt(8)

    p_link = doc.add_paragraph()
    p_link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(p_link, APP_URL, APP_URL)
    p_link.paragraph_format.space_after = Pt(6)

    p_tip = doc.add_paragraph()
    run_tip = p_tip.add_run("ブックマーク推奨: ")
    run_tip.bold = True
    run_tip.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    p_tip.add_run("上記URLをブラウザのブックマークに登録しておくと、次回からすぐアクセスできます。")
    p_tip.paragraph_format.space_after = Pt(16)

    # =============================================
    # 1. このツールについて
    # =============================================
    doc.add_heading("1. このツールについて", level=1)

    doc.add_paragraph(
        "QPSS 脚本フォーマッタは、脚本・シナリオのファイルを"
        "QPSS準拠のフォーマット（フォント・縦書き・インデント等）に整形し、"
        "Word文書（.docx）としてダウンロードできるWebツールです。"
    )

    doc.add_heading("特徴", level=3)
    features = [
        "ブラウザだけで使える（インストール不要）",
        "ファイルをアップロードするだけで、フォーマット済みWord文書を生成",
        ".txt / .docx / .rtf / .md の4形式に対応",
        "プリセットで案件ごとにフォーマットを切替",
        "ト書きの自動インデント",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Bullet")

    # =============================================
    # 2. 基本的な使い方
    # =============================================
    doc.add_heading("2. 基本的な使い方", level=1)

    # Step 1
    doc.add_heading("Step 1: アプリを開く", level=2)
    p = doc.add_paragraph("上記のリンクをクリックしてアプリを開きます。")
    p.paragraph_format.space_after = Pt(8)

    # Step 2
    doc.add_heading("Step 2: 脚本ファイルをアップロード", level=2)
    doc.add_paragraph(
        "画面の「脚本ファイルをアップロード」エリアに、"
        "ファイルをドラッグ＆ドロップするか、「Browse files」ボタンで選択します。"
    )

    doc.add_heading("対応ファイル形式", level=3)
    t = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ("形式", "説明"),
        (".txt", "メモ帳などで作成したテキストファイル"),
        (".docx", "Microsoft Word文書"),
        (".rtf", "リッチテキスト形式"),
        (".md", "Markdownファイル"),
    ]
    for i, (c1, c2) in enumerate(data):
        t.rows[i].cells[0].text = c1
        t.rows[i].cells[1].text = c2
        if i == 0:
            for cell in t.rows[0].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
    add_styled_paragraph(doc, "", space_after=6)

    # Step 3
    doc.add_heading("Step 3: 執筆者名と出力ファイル名を入力", level=2)
    items = [
        "「執筆者名」に名前を入力すると、タイトルの次の行に「脚本：○○」として挿入されます。空欄でもOKです。",
        "「出力ファイル名」でダウンロードするファイルの名前を変更できます。",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    # Step 4
    doc.add_heading("Step 4: プリセットを選択", level=2)
    doc.add_paragraph("「プリセットを選択」のドロップダウンから、フォーマット設定を選びます。")

    pt = doc.add_table(rows=3, cols=3, style="Light Grid Accent 1")
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER
    pdata = [
        ("プリセット名", "ページの向き", "文字方向"),
        ("QPSSフォーマッタ標準", "横向き", "縦書き"),
        ("横書き標準", "縦向き", "横書き"),
    ]
    for i, row in enumerate(pdata):
        for j, val in enumerate(row):
            pt.rows[i].cells[j].text = val
            if i == 0:
                for p in pt.rows[0].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True
    add_styled_paragraph(doc, "", space_after=6)

    # Step 5
    doc.add_heading("Step 5: Word文書を生成・ダウンロード", level=2)
    steps = [
        "「Word文書を生成」ボタンをクリック",
        "「Word文書の生成が完了しました」と表示されたら成功",
        "「ダウンロード」ボタンをクリックしてファイルを保存",
    ]
    for idx, s in enumerate(steps, 1):
        doc.add_paragraph(f"{idx}. {s}")

    # =============================================
    # 3. フォーマットのカスタマイズ
    # =============================================
    doc.add_heading("3. フォーマットのカスタマイズ", level=1)

    doc.add_paragraph(
        "「詳細設定（クリックで展開）」をクリックすると、フォーマットの詳細を変更できます。"
    )

    doc.add_heading("ページ・フォント設定", level=2)
    st = doc.add_table(rows=8, cols=3, style="Light Grid Accent 1")
    st.alignment = WD_TABLE_ALIGNMENT.CENTER
    sdata = [
        ("設定項目", "初期値（QPSSフォーマッタ標準）", "説明"),
        ("ページの向き", "横向き", "横向き / 縦向き"),
        ("文字方向", "縦書き", "縦書き / 横書き"),
        ("用紙サイズ", "A4", "A4 / B5 / Letter"),
        ("フォント", "MS Gothic", "フォント名を入力"),
        ("フォントサイズ", "12 pt", "6〜72 ptで指定"),
        ("前インデント", "6 mm", "文字の前（上）の余白"),
        ("ページ番号", "あり（右下）", "表示位置も変更可能"),
    ]
    for i, row in enumerate(sdata):
        for j, val in enumerate(row):
            st.rows[i].cells[j].text = val
            if i == 0:
                for p in st.rows[0].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True
    add_styled_paragraph(doc, "", space_after=6)

    doc.add_heading("ト書き判定ルール", level=2)
    doc.add_paragraph(
        "ト書き（登場人物の動作・情景描写）は自動判定され、TABインデントが付与されます。"
        "セリフ・柱・テロップ等は「除外パターン」で定義されています。"
    )
    doc.add_paragraph(
        "案件ごとに固有のタイトルや名前がある場合は、「詳細設定」→「ト書き判定」タブから"
        "除外パターンを追加・削除できます。"
    )

    doc.add_heading("パターンタイプの説明", level=3)
    tt = doc.add_table(rows=5, cols=3, style="Light Grid Accent 1")
    tt.alignment = WD_TABLE_ALIGNMENT.CENTER
    tdata = [
        ("タイプ", "意味", "例"),
        ("contains", "指定文字列を含む行", "「 → セリフ行を除外"),
        ("starts_with", "指定文字列で始まる行", "脚本 → 脚本表記行を除外"),
        ("exact", "完全一致する行", "＜了＞ → 終了記号を除外"),
        ("regex", "正規表現でマッチする行", "^[０-９] → 柱行を除外"),
    ]
    for i, row in enumerate(tdata):
        for j, val in enumerate(row):
            tt.rows[i].cells[j].text = val
            if i == 0:
                for p in tt.rows[0].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True
    add_styled_paragraph(doc, "", space_after=6)

    # =============================================
    # 4. 設定のダウンロード
    # =============================================
    doc.add_heading("4. 設定のバックアップ", level=1)

    doc.add_paragraph(
        "詳細設定で変更した内容は、YAMLファイルとしてダウンロードし、手元に保存しておくことができます。"
        "次回同じ設定を使いたいときに参照用として活用できます。"
    )
    save_steps = [
        "「詳細設定」を展開し、フォーマットを好みに調整",
        "一番下の入力欄にプリセット名を入力（例: 案件A用）",
        "「保存」ボタンをクリック → YAMLファイルがダウンロードされます",
    ]
    for idx, s in enumerate(save_steps, 1):
        doc.add_paragraph(f"{idx}. {s}")

    p_note = doc.add_paragraph()
    run_note = p_note.add_run("※ ")
    run_note.bold = True
    p_note.add_run(
        "ダウンロードしたYAMLは設定の記録・バックアップ用です。"
        "次回利用時は、詳細設定から同じ値を再入力してください。"
    )

    # =============================================
    # 5. よくある質問
    # =============================================
    doc.add_heading("5. よくある質問", level=1)

    faq = doc.add_table(rows=6, cols=2, style="Light Grid Accent 1")
    faq.alignment = WD_TABLE_ALIGNMENT.CENTER
    faq_data = [
        ("質問", "回答"),
        ("インストールは必要？",
         "不要です。ブラウザでURLにアクセスするだけで使えます。"),
        ("スマホでも使える？",
         "ブラウザがあれば使えますが、PCでの利用を推奨します。"),
        ("アップロードしたファイルは\n保存される？",
         "保存されません。ファイルはサーバー上で一時的に処理され、\nセッション終了後に自動削除されます。"),
        ("文字化けする場合は？",
         "入力ファイルの文字コードをUTF-8で保存し直してください。\nメモ帳: 名前を付けて保存 → 文字コード → UTF-8"),
        ("ト書き判定がおかしい場合は？",
         "「詳細設定」→「ト書き判定」タブで\n除外パターンを追加・調整してください。"),
    ]
    for i, (c1, c2) in enumerate(faq_data):
        faq.rows[i].cells[0].text = c1
        faq.rows[i].cells[1].text = c2
        if i == 0:
            for cell in faq.rows[0].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    # =============================================
    # 6. 情報セキュリティについて
    # =============================================
    doc.add_heading("6. 情報セキュリティについて", level=1)

    doc.add_heading("安心して利用できる点", level=2)
    safe_points = [
        "アップロードしたファイルはサーバーに保存されません。セッション中のみメモリ上で処理され、終了後に自動削除されます。",
        "ファイルの変換処理はサーバー内で完結し、外部APIへの転送は一切行いません。",
        "GoogleアカウントやSNS等への接続・認証は不要です。",
        "ブラウザとサーバー間の通信はHTTPS（SSL/TLS）で暗号化されています。",
        "ソースコードはGitHubで公開されており、不審な処理がないことを誰でも検証できます。",
    ]
    for s in safe_points:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("注意が必要な点", level=2)
    risk_table = doc.add_table(rows=5, cols=3, style="Light Grid Accent 1")
    risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    risk_data = [
        ("リスク", "詳細", "対策"),
        ("アプリURLが公開状態",
         "URLを知っていれば誰でもアクセス可能",
         "URLの社外共有を避ける"),
        ("脚本内容がインターネットを経由",
         "ファイルアップロード時にクラウドサーバー\n（米国）を経由する",
         "機密性の高い脚本はローカル版\n（start.bat）を使用する"),
        ("プリセットYAMLがGitHubで公開",
         "案件固有の情報をプリセットに書くと\n外部から見える",
         "プリセットに案件名・クライアント名を\n含めない"),
        ("Streamlit Cloudの利用規約",
         "Snowflake社のサービスであり、\n利用規約に準拠する",
         "社内のクラウドサービス利用ポリシーを\n確認する"),
    ]
    for i, row in enumerate(risk_data):
        for j, val in enumerate(row):
            risk_table.rows[i].cells[j].text = val
            if i == 0:
                for cell in risk_table.rows[0].cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
    add_styled_paragraph(doc, "", space_after=6)

    doc.add_heading("推奨運用ルール", level=2)
    rules = [
        "機密性の高い脚本 → ローカル版（start.bat）を使用（インターネット不使用）",
        "一般的な脚本 → Web版（このアプリ）でOK",
        "プリセットYAML → 案件名・クライアント名を含めない",
        "アプリURL → 社内のみで共有し、外部に公開しない",
    ]
    for idx, rule in enumerate(rules, 1):
        doc.add_paragraph(f"{idx}. {rule}")

    add_styled_paragraph(doc, "", space_after=20)

    # フッター
    add_styled_paragraph(
        doc, "QPSS 脚本フォーマッタ 使い方ガイド（Web版）  |  Alche",
        font_size=9, color=(0xAA, 0xAA, 0xAA),
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )

    # 保存
    output_path = Path(__file__).parent.parent / "QPSS脚本フォーマッタ_使い方ガイド_Web版.docx"
    doc.save(str(output_path))
    print(f"ガイド生成完了: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    create_guide()
